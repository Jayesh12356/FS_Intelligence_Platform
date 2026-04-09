"""Export API — JIRA, Confluence, PDF, DOCX, test-cases (L10).

Endpoints:
  POST /api/fs/{id}/export/jira       — export tasks to JIRA
  POST /api/fs/{id}/export/confluence  — export analysis to Confluence
  GET  /api/fs/{id}/export/pdf        — download PDF report
  GET  /api/fs/{id}/export/docx       — download Word report
  GET  /api/fs/{id}/test-cases        — list generated test cases
  GET  /api/fs/{id}/test-cases/csv    — export test cases as CSV
"""

import csv
import io
import logging
import uuid

from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import StreamingResponse
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.db.base import get_db
from app.db.models import (
    AmbiguityFlagDB,
    ComplianceTagDB,
    ContradictionDB,
    EdgeCaseGapDB,
    FSDocument,
    FSDocumentStatus,
    FSTaskDB,
    TestCaseDB,
    TraceabilityEntryDB,
)
from app.models.schemas import (
    APIResponse,
    ConfluenceExportResponse,
    JiraExportResponse,
    ReportExportResponse,
    TestCaseListResponse,
    TestCaseSchema,
)

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/api/fs", tags=["export"])


# ── Helper: Fetch document or 404 ─────────────────────


async def _get_doc(doc_id: uuid.UUID, db: AsyncSession) -> FSDocument:
    result = await db.execute(select(FSDocument).where(FSDocument.id == doc_id))
    doc = result.scalar_one_or_none()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    return doc


# ── JIRA Export ────────────────────────────────────────


@router.post("/{doc_id}/export/jira", response_model=APIResponse[JiraExportResponse])
async def export_to_jira(
    doc_id: uuid.UUID,
    db: AsyncSession = Depends(get_db),
) -> APIResponse[JiraExportResponse]:
    """Export FS tasks to JIRA as an epic with stories."""
    doc = await _get_doc(doc_id, db)

    # Load tasks
    result = await db.execute(
        select(FSTaskDB).where(FSTaskDB.fs_id == doc_id).order_by(FSTaskDB.order)
    )
    tasks = result.scalars().all()

    if not tasks:
        raise HTTPException(status_code=400, detail="No tasks to export — run analysis first")

    task_dicts = [
        {
            "task_id": t.task_id,
            "title": t.title,
            "description": t.description,
            "section_heading": t.section_heading,
            "acceptance_criteria": t.acceptance_criteria or [],
            "effort": t.effort.value if t.effort else "MEDIUM",
            "tags": t.tags or [],
        }
        for t in tasks
    ]

    from app.integrations.jira import JiraClient

    client = JiraClient()
    export_result = await client.export_fs_tasks(doc.filename, task_dicts)

    # Log audit event
    from app.db.audit import log_audit_event
    from app.db.models import AuditEventType
    await log_audit_event(
        db, doc_id, AuditEventType.EXPORTED,
        payload={"target": "jira", "stories": len(export_result["stories"])},
    )
    await db.commit()

    return APIResponse(
        data=JiraExportResponse(
            epic=export_result["epic"],
            stories=export_result["stories"],
            total=export_result["total"],
            simulated=export_result["epic"].get("simulated", False),
        ),
    )


# ── Confluence Export ──────────────────────────────────


@router.post("/{doc_id}/export/confluence", response_model=APIResponse[ConfluenceExportResponse])
async def export_to_confluence(
    doc_id: uuid.UUID,
    db: AsyncSession = Depends(get_db),
) -> APIResponse[ConfluenceExportResponse]:
    """Export full FS analysis to a Confluence page."""
    doc = await _get_doc(doc_id, db)

    # Load all analysis data
    sections = []
    if doc.parsed_text:
        from app.parsers.chunker import chunk_text_into_sections
        sections = chunk_text_into_sections(doc.parsed_text)

    # Load tasks
    tasks_result = await db.execute(
        select(FSTaskDB).where(FSTaskDB.fs_id == doc_id).order_by(FSTaskDB.order)
    )
    tasks = [
        {
            "task_id": t.task_id,
            "title": t.title,
            "effort": t.effort.value if t.effort else "MEDIUM",
            "section_heading": t.section_heading,
        }
        for t in tasks_result.scalars().all()
    ]

    # Load ambiguities
    amb_result = await db.execute(
        select(AmbiguityFlagDB).where(AmbiguityFlagDB.fs_id == doc_id)
    )
    ambiguities = [
        {
            "section_heading": a.section_heading,
            "severity": a.severity.value if a.severity else "MEDIUM",
            "reason": a.reason,
        }
        for a in amb_result.scalars().all()
    ]

    # Load traceability
    trace_result = await db.execute(
        select(TraceabilityEntryDB).where(TraceabilityEntryDB.fs_id == doc_id)
    )
    traceability = [
        {
            "task_id": t.task_id,
            "task_title": t.task_title,
            "section_heading": t.section_heading,
        }
        for t in trace_result.scalars().all()
    ]

    from app.integrations.confluence import ConfluenceClient

    client = ConfluenceClient()
    page_result = await client.create_fs_page(
        title=doc.filename,
        sections=sections,
        quality_score=None,  # TODO: load from analysis if stored
        ambiguities=ambiguities,
        tasks=tasks,
        traceability=traceability,
    )

    # Log audit
    from app.db.audit import log_audit_event
    from app.db.models import AuditEventType
    await log_audit_event(
        db, doc_id, AuditEventType.EXPORTED,
        payload={"target": "confluence", "page_id": page_result.get("id", "")},
    )
    await db.commit()

    return APIResponse(
        data=ConfluenceExportResponse(
            page_id=page_result["id"],
            page_url=page_result["url"],
            title=page_result["title"],
            simulated=page_result.get("simulated", False),
        ),
    )


# ── Test Cases ─────────────────────────────────────────


@router.get("/{doc_id}/test-cases", response_model=APIResponse[TestCaseListResponse])
async def list_test_cases(
    doc_id: uuid.UUID,
    db: AsyncSession = Depends(get_db),
) -> APIResponse[TestCaseListResponse]:
    """List all generated test cases for a document."""
    await _get_doc(doc_id, db)

    result = await db.execute(
        select(TestCaseDB)
        .where(TestCaseDB.fs_id == doc_id)
        .order_by(TestCaseDB.section_index, TestCaseDB.task_id)
    )
    rows = result.scalars().all()

    schemas = [
        TestCaseSchema(
            id=r.id,
            fs_id=r.fs_id,
            task_id=r.task_id,
            title=r.title,
            preconditions=r.preconditions or "",
            steps=r.steps or [],
            expected_result=r.expected_result or "",
            test_type=r.test_type.value if r.test_type else "UNIT",
            section_index=r.section_index,
            section_heading=r.section_heading or "",
            created_at=r.created_at,
        )
        for r in rows
    ]

    # Count by type
    by_type: dict = {}
    for tc in schemas:
        by_type[tc.test_type] = by_type.get(tc.test_type, 0) + 1

    return APIResponse(
        data=TestCaseListResponse(
            test_cases=schemas,
            total=len(schemas),
            by_type=by_type,
        ),
    )


@router.get("/{doc_id}/test-cases/csv")
async def export_test_cases_csv(
    doc_id: uuid.UUID,
    db: AsyncSession = Depends(get_db),
):
    """Export test cases as CSV."""
    await _get_doc(doc_id, db)

    result = await db.execute(
        select(TestCaseDB)
        .where(TestCaseDB.fs_id == doc_id)
        .order_by(TestCaseDB.section_index, TestCaseDB.task_id)
    )
    rows = result.scalars().all()

    output = io.StringIO()
    writer = csv.writer(output)
    writer.writerow(["Task ID", "Title", "Type", "Preconditions", "Steps", "Expected Result", "Section"])

    for r in rows:
        steps_str = "; ".join(r.steps or [])
        writer.writerow([
            r.task_id,
            r.title,
            r.test_type.value if r.test_type else "UNIT",
            r.preconditions or "",
            steps_str,
            r.expected_result or "",
            r.section_heading or "",
        ])

    output.seek(0)
    return StreamingResponse(
        iter([output.getvalue()]),
        media_type="text/csv",
        headers={"Content-Disposition": f"attachment; filename=test_cases_{doc_id}.csv"},
    )


# ── PDF Report Export ──────────────────────────────────


@router.get("/{doc_id}/export/pdf", response_model=APIResponse[ReportExportResponse])
async def export_pdf_report(
    doc_id: uuid.UUID,
    db: AsyncSession = Depends(get_db),
) -> APIResponse[ReportExportResponse]:
    """Generate and return a PDF intelligence report.

    Uses basic HTML-to-PDF approach since reportlab may not be installed.
    Returns download metadata with the report content.
    """
    doc = await _get_doc(doc_id, db)

    # Build report content
    report_content = await _build_report_content(doc_id, doc, db)

    # Generate PDF
    try:
        pdf_bytes = _generate_pdf(doc.filename, report_content)
    except Exception as exc:
        logger.error("PDF generation failed: %s", exc)
        raise HTTPException(status_code=500, detail=f"PDF generation failed: {exc}")

    filename = f"FS_Report_{doc.filename.replace('.', '_')}.pdf"

    # Log audit
    from app.db.audit import log_audit_event
    from app.db.models import AuditEventType
    await log_audit_event(
        db, doc_id, AuditEventType.EXPORTED,
        payload={"target": "pdf", "filename": filename},
    )
    await db.commit()

    return APIResponse(
        data=ReportExportResponse(
            filename=filename,
            format="pdf",
            size_bytes=len(pdf_bytes),
            download_url=f"/api/fs/{doc_id}/export/pdf/download",
        ),
    )


@router.get("/{doc_id}/export/pdf/download")
async def download_pdf_report(
    doc_id: uuid.UUID,
    db: AsyncSession = Depends(get_db),
):
    """Download the actual PDF file."""
    doc = await _get_doc(doc_id, db)
    report_content = await _build_report_content(doc_id, doc, db)
    pdf_bytes = _generate_pdf(doc.filename, report_content)
    filename = f"FS_Report_{doc.filename.replace('.', '_')}.pdf"

    return StreamingResponse(
        iter([pdf_bytes]),
        media_type="application/pdf",
        headers={"Content-Disposition": f"attachment; filename={filename}"},
    )


# ── DOCX Report Export ─────────────────────────────────


@router.get("/{doc_id}/export/docx", response_model=APIResponse[ReportExportResponse])
async def export_docx_report(
    doc_id: uuid.UUID,
    db: AsyncSession = Depends(get_db),
) -> APIResponse[ReportExportResponse]:
    """Generate and return a Word (.docx) intelligence report."""
    doc = await _get_doc(doc_id, db)
    report_content = await _build_report_content(doc_id, doc, db)

    try:
        docx_bytes = _generate_docx(doc.filename, report_content)
    except Exception as exc:
        logger.error("DOCX generation failed: %s", exc)
        raise HTTPException(status_code=500, detail=f"DOCX generation failed: {exc}")

    filename = f"FS_Report_{doc.filename.replace('.', '_')}.docx"

    from app.db.audit import log_audit_event
    from app.db.models import AuditEventType
    await log_audit_event(
        db, doc_id, AuditEventType.EXPORTED,
        payload={"target": "docx", "filename": filename},
    )
    await db.commit()

    return APIResponse(
        data=ReportExportResponse(
            filename=filename,
            format="docx",
            size_bytes=len(docx_bytes),
            download_url=f"/api/fs/{doc_id}/export/docx/download",
        ),
    )


@router.get("/{doc_id}/export/docx/download")
async def download_docx_report(
    doc_id: uuid.UUID,
    db: AsyncSession = Depends(get_db),
):
    """Download the actual DOCX file."""
    doc = await _get_doc(doc_id, db)
    report_content = await _build_report_content(doc_id, doc, db)
    docx_bytes = _generate_docx(doc.filename, report_content)
    filename = f"FS_Report_{doc.filename.replace('.', '_')}.docx"

    return StreamingResponse(
        iter([docx_bytes]),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={filename}"},
    )


# ── Report Helpers ─────────────────────────────────────


async def _build_report_content(
    doc_id: uuid.UUID,
    doc: FSDocument,
    db: AsyncSession,
) -> dict:
    """Gather all analysis data for report generation."""
    # Tasks
    tasks_result = await db.execute(
        select(FSTaskDB).where(FSTaskDB.fs_id == doc_id).order_by(FSTaskDB.order)
    )
    tasks = tasks_result.scalars().all()

    # Ambiguities
    amb_result = await db.execute(
        select(AmbiguityFlagDB).where(AmbiguityFlagDB.fs_id == doc_id)
    )
    ambiguities = amb_result.scalars().all()

    # Contradictions
    contra_result = await db.execute(
        select(ContradictionDB).where(ContradictionDB.fs_id == doc_id)
    )
    contradictions = contra_result.scalars().all()

    # Edge cases
    ec_result = await db.execute(
        select(EdgeCaseGapDB).where(EdgeCaseGapDB.fs_id == doc_id)
    )
    edge_cases = ec_result.scalars().all()

    # Compliance
    comp_result = await db.execute(
        select(ComplianceTagDB).where(ComplianceTagDB.fs_id == doc_id)
    )
    compliance = comp_result.scalars().all()

    # Traceability
    trace_result = await db.execute(
        select(TraceabilityEntryDB).where(TraceabilityEntryDB.fs_id == doc_id)
    )
    traceability = trace_result.scalars().all()

    # Test cases
    tc_result = await db.execute(
        select(TestCaseDB).where(TestCaseDB.fs_id == doc_id)
    )
    test_cases = tc_result.scalars().all()

    return {
        "filename": doc.filename,
        "status": doc.status.value if doc.status else "UNKNOWN",
        "tasks": tasks,
        "ambiguities": ambiguities,
        "contradictions": contradictions,
        "edge_cases": edge_cases,
        "compliance": compliance,
        "traceability": traceability,
        "test_cases": test_cases,
    }


def _generate_pdf(filename: str, content: dict) -> bytes:
    """Generate a simple PDF report from analysis content.

    Uses basic byte construction since reportlab might not be available.
    """
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.lib import colors
        import io as _io

        buffer = _io.BytesIO()
        doc_pdf = SimpleDocTemplate(buffer, pagesize=A4)
        styles = getSampleStyleSheet()
        story = []

        # Title
        story.append(Paragraph(f"FS Intelligence Report: {filename}", styles["Title"]))
        story.append(Spacer(1, 12))
        story.append(Paragraph(f"Status: {content['status']}", styles["Normal"]))
        story.append(Spacer(1, 24))

        # Summary stats
        stats = [
            ["Metric", "Count"],
            ["Tasks", str(len(content["tasks"]))],
            ["Ambiguity Flags", str(len(content["ambiguities"]))],
            ["Contradictions", str(len(content["contradictions"]))],
            ["Edge Cases", str(len(content["edge_cases"]))],
            ["Compliance Tags", str(len(content["compliance"]))],
            ["Test Cases", str(len(content["test_cases"]))],
        ]
        t = Table(stats)
        t.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#8b5cf6")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("ALIGN", (0, 0), (-1, -1), "LEFT"),
            ("FONTSIZE", (0, 0), (-1, -1), 10),
            ("BOTTOMPADDING", (0, 0), (-1, 0), 12),
            ("BACKGROUND", (0, 1), (-1, -1), colors.HexColor("#f5f3ff")),
            ("GRID", (0, 0), (-1, -1), 1, colors.HexColor("#c4b5fd")),
        ]))
        story.append(t)
        story.append(Spacer(1, 20))

        # Tasks
        if content["tasks"]:
            story.append(Paragraph("Task Breakdown", styles["Heading2"]))
            for task in content["tasks"]:
                story.append(Paragraph(
                    f"<b>{task.task_id}</b>: {task.title} (Effort: {task.effort.value})",
                    styles["Normal"],
                ))
            story.append(Spacer(1, 12))

        # Ambiguities
        if content["ambiguities"]:
            story.append(Paragraph("Ambiguity Flags", styles["Heading2"]))
            for amb in content["ambiguities"]:
                story.append(Paragraph(
                    f"[{amb.severity.value}] {amb.section_heading}: {amb.reason}",
                    styles["Normal"],
                ))
            story.append(Spacer(1, 12))

        # Test Cases
        if content["test_cases"]:
            story.append(Paragraph("Test Cases", styles["Heading2"]))
            for tc in content["test_cases"]:
                story.append(Paragraph(
                    f"<b>{tc.title}</b> ({tc.test_type.value}) — Task: {tc.task_id}",
                    styles["Normal"],
                ))

        doc_pdf.build(story)
        return buffer.getvalue()

    except ImportError:
        # Fallback: generate a simple text-based pseudo-PDF
        logger.warning("reportlab not installed — generating simple text report")
        return _generate_text_report(filename, content, "PDF")


def _generate_docx(filename: str, content: dict) -> bytes:
    """Generate a Word .docx report from analysis content."""
    try:
        from docx import Document as DocxDocument
        from docx.shared import Inches, Pt, RGBColor
        import io as _io

        doc = DocxDocument()

        # Title
        title = doc.add_heading(f"FS Intelligence Report: {filename}", 0)

        # Status
        doc.add_paragraph(f"Status: {content['status']}")
        doc.add_paragraph("")

        # Summary table
        doc.add_heading("Summary", level=1)
        table = doc.add_table(rows=1, cols=2)
        table.style = "Light Grid Accent 1"
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Metric"
        hdr_cells[1].text = "Count"

        for metric, count in [
            ("Tasks", len(content["tasks"])),
            ("Ambiguity Flags", len(content["ambiguities"])),
            ("Contradictions", len(content["contradictions"])),
            ("Edge Cases", len(content["edge_cases"])),
            ("Compliance Tags", len(content["compliance"])),
            ("Test Cases", len(content["test_cases"])),
        ]:
            row = table.add_row().cells
            row[0].text = metric
            row[1].text = str(count)

        # Tasks
        if content["tasks"]:
            doc.add_heading("Task Breakdown", level=1)
            for task in content["tasks"]:
                doc.add_paragraph(
                    f"{task.task_id}: {task.title} (Effort: {task.effort.value})",
                    style="List Bullet",
                )

        # Ambiguities
        if content["ambiguities"]:
            doc.add_heading("Ambiguity Flags", level=1)
            for amb in content["ambiguities"]:
                doc.add_paragraph(
                    f"[{amb.severity.value}] {amb.section_heading}: {amb.reason}",
                    style="List Bullet",
                )

        # Test Cases
        if content["test_cases"]:
            doc.add_heading("Test Cases", level=1)
            tc_table = doc.add_table(rows=1, cols=4)
            tc_table.style = "Light Grid Accent 1"
            hdr = tc_table.rows[0].cells
            hdr[0].text = "Task ID"
            hdr[1].text = "Title"
            hdr[2].text = "Type"
            hdr[3].text = "Expected Result"
            for tc in content["test_cases"]:
                row = tc_table.add_row().cells
                row[0].text = tc.task_id
                row[1].text = tc.title
                row[2].text = tc.test_type.value if tc.test_type else "UNIT"
                row[3].text = tc.expected_result or ""

        # Traceability
        if content["traceability"]:
            doc.add_heading("Traceability Matrix", level=1)
            tr_table = doc.add_table(rows=1, cols=3)
            tr_table.style = "Light Grid Accent 1"
            hdr = tr_table.rows[0].cells
            hdr[0].text = "Task ID"
            hdr[1].text = "Task Title"
            hdr[2].text = "Section"
            for tr in content["traceability"]:
                row = tr_table.add_row().cells
                row[0].text = tr.task_id
                row[1].text = tr.task_title
                row[2].text = tr.section_heading

        buffer = _io.BytesIO()
        doc.save(buffer)
        return buffer.getvalue()

    except ImportError:
        logger.warning("python-docx not installed — generating simple text report")
        return _generate_text_report(filename, content, "DOCX")


def _generate_text_report(filename: str, content: dict, fmt: str) -> bytes:
    """Fallback text report when reportlab/python-docx is unavailable."""
    lines = [
        f"FS Intelligence Report: {filename}",
        f"Format: {fmt}",
        f"Status: {content['status']}",
        "",
        "=" * 60,
        "SUMMARY",
        "=" * 60,
        f"Tasks: {len(content['tasks'])}",
        f"Ambiguity Flags: {len(content['ambiguities'])}",
        f"Contradictions: {len(content['contradictions'])}",
        f"Edge Cases: {len(content['edge_cases'])}",
        f"Compliance Tags: {len(content['compliance'])}",
        f"Test Cases: {len(content['test_cases'])}",
        "",
    ]

    if content["tasks"]:
        lines.append("=" * 60)
        lines.append("TASKS")
        lines.append("=" * 60)
        for t in content["tasks"]:
            lines.append(f"  {t.task_id}: {t.title} [Effort: {t.effort.value}]")
        lines.append("")

    if content["ambiguities"]:
        lines.append("=" * 60)
        lines.append("AMBIGUITY FLAGS")
        lines.append("=" * 60)
        for a in content["ambiguities"]:
            lines.append(f"  [{a.severity.value}] {a.section_heading}: {a.reason}")
        lines.append("")

    if content["test_cases"]:
        lines.append("=" * 60)
        lines.append("TEST CASES")
        lines.append("=" * 60)
        for tc in content["test_cases"]:
            tc_type = tc.test_type.value if tc.test_type else "UNIT"
            lines.append(f"  [{tc_type}] {tc.title} — Task: {tc.task_id}")
        lines.append("")

    return "\n".join(lines).encode("utf-8")
