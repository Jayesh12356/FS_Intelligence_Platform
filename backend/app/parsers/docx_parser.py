"""DOCX parser — extracts structured sections from Word documents.

Uses python-docx to read paragraph styles and map Word heading levels
to FSSection instances. Preserves numbered lists as requirement candidates.
"""

import logging
from pathlib import Path
from typing import List

from docx import Document as DocxDocument

from app.parsers.base import FSSection, ParsedFS

logger = logging.getLogger(__name__)


def parse_docx(filepath: str) -> ParsedFS:
    """Parse a DOCX file into structured sections.

    Maps Word heading styles (Heading 1–4) to section headings.
    Regular paragraphs are grouped under the most recent heading.

    Args:
        filepath: Path to the DOCX file on disk.

    Returns:
        ParsedFS with raw text and heading-based sections.

    Raises:
        FileNotFoundError: If the file doesn't exist.
        ValueError: If the DOCX is corrupted or unreadable.
    """
    path = Path(filepath)
    if not path.exists():
        raise FileNotFoundError(f"DOCX file not found: {filepath}")

    try:
        doc = DocxDocument(str(path))
    except Exception as exc:
        logger.error("Failed to open DOCX %s: %s", filepath, exc)
        raise ValueError(f"Cannot read DOCX file: {exc}") from exc

    sections: List[FSSection] = []
    current_heading = "Introduction"
    current_content: List[str] = []
    section_index = 0
    all_text_parts: List[str] = []

    heading_styles = {
        "Heading 1",
        "Heading 2",
        "Heading 3",
        "Heading 4",
        "heading 1",
        "heading 2",
        "heading 3",
        "heading 4",
        "Title",
        "Subtitle",
    }

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        all_text_parts.append(text)
        style_name = para.style.name if para.style else ""

        if style_name in heading_styles:
            # Save previous section
            content_text = "\n".join(current_content).strip()
            if content_text:
                sections.append(
                    FSSection(
                        heading=current_heading,
                        content=content_text,
                        section_index=section_index,
                    )
                )
                section_index += 1
            current_heading = text
            current_content = []
        else:
            # Format list items with bullets for readability
            if style_name and "list" in style_name.lower():
                text = f"• {text}"
            current_content.append(text)

    # Save the last section
    content_text = "\n".join(current_content).strip()
    if content_text:
        sections.append(
            FSSection(
                heading=current_heading,
                content=content_text,
                section_index=section_index,
            )
        )

    raw_text = "\n\n".join(all_text_parts)

    # Fallback: if no headings detected, wrap all content in one section
    if not sections and raw_text.strip():
        sections.append(
            FSSection(
                heading="Document Content",
                content=raw_text.strip(),
                section_index=0,
            )
        )

    logger.info(
        "Parsed DOCX %s: %d paragraphs, %d sections, %d chars",
        filepath,
        len(doc.paragraphs),
        len(sections),
        len(raw_text),
    )

    return ParsedFS(
        raw_text=raw_text,
        sections=sections,
        metadata={
            "paragraphs": len(doc.paragraphs),
            "parser": "python-docx",
            "characters": len(raw_text),
        },
    )
